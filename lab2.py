from openpyxl import Workbook
from openpyxl.styles import PatternFill
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import os



# Вариант В, Б, Г, А, В 
# 1) Источник данных - Прочитать данные из JSON-файла (формат: список словарей с ключами имя,количество,цена).
# 2) Операция с данными в Excel - Добавить условное форматирование: выделить ячейки в колонке "Количество" цветом, если значение больше 100.
# 3) Структура Word-отчета - Создать документ с двумя разделами: первый - исходные данные (таблица), второй - итоги (итоговая сумма и самая дорогая позиция).
# 4) Действие с PDF - Конвертировать готовый Word-документ в PDF.
# 5) Финализация работы - Закрыть все открытые приложения (если использовались) и вывести в консоль путь к созданным файлам.

def parser(filename: str) -> dict:
    with open(filename, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data


def createExcel(filename: str, data: dict):
    excel = Workbook()
    sheet = excel.active

    sheet['A1'] = 'Ключ'
    sheet['B1'] = 'Наименование'
    sheet['C1'] = 'Количество'
    sheet['D1'] = 'Цена'

    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    currow = 2

    for product in data["products"]:
        sheet.cell(row=currow, column=1, value=product["key"])
        sheet.cell(row=currow, column=2, value=product["name"])
        sheet.cell(row=currow, column=3, value=product["count"])
        sheet.cell(row=currow, column=4, value=product["price"])
        if product["count"] > 100:
            sheet['C' + str(currow)].fill = yellow_fill
        currow += 1

    excel.save(filename + '.xlsx')

def createWordAndPDF(filename: str, data: dict):
    doc = Document()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ключ'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Количество'
    hdr_cells[3].text = 'Цена'

    sum_cost = 0
    max_cost_key = 0
    max_cost = 0

    for product in data["products"]:
        row_cells = table.add_row().cells
        row_cells[0].text = str(product["key"])
        row_cells[1].text = str(product["name"])
        row_cells[2].text = str(product["count"])
        row_cells[3].text = str(product["price"])

        cost = product["count"] * product["price"]
        if cost > max_cost:
            max_cost = cost
            max_cost_key = product["key"]
        sum_cost += cost


    doc.add_paragraph().add_run('Итоги:').bold = True
    doc.add_paragraph().add_run('Итоговая сумма: ' + str(sum_cost))
    doc.add_paragraph().add_run('Самая дорога позиция: ' + str(max_cost_key) + ' Её стоимость: ' + str(max_cost))
    
    doc.save(filename + '.docx')
    
    convert(filename + '.docx', filename + '.pdf')

def get_file_paths(filename):
    abs_path = os.path.abspath(filename)
    return abs_path



data = parser('lab2_products.json')

createExcel('Отчёт', data)
createWordAndPDF('Отчёт', data)
print(get_file_paths('Отчёт.xlsx'))
print(get_file_paths('Отчёт.docx'))
print(get_file_paths('Отчёт.pdf'))
